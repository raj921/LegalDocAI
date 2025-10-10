from docx import Document
import pdfplumber
import logging
from typing import List, Tuple

logger = logging.getLogger(__name__)

class DocumentService:
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return text
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            raise
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            raise
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
        
        return chunks if chunks else [text]
    
    @staticmethod
    def create_markdown_template(title: str, content: str, variables: List[dict], 
                                 doc_type: str, jurisdiction: str, tags: List[str],
                                 template_id: str, description: str = "") -> str:
        """Create a markdown template with YAML front-matter and variable placeholders."""
        import re
        
        # YAML front-matter with proper structure
        front_matter = f"""---
template_id: {template_id}
title: {title}
File_description: {description or f'Template for {doc_type}'}
jurisdiction: {jurisdiction or 'N/A'}
doc_type: {doc_type}
variables:\n"""
        
        # Add variables in YAML format
        if variables:
            for var in variables:
                key = var.get('key', 'unknown')
                label = var.get('label', key.replace('_', ' ').title())
                desc = var.get('description', '')
                example = var.get('example', '')
                required = var.get('is_required', True)
                
                front_matter += f"  - key: {key}\n"
                front_matter += f"    label: {label}\n"
                front_matter += f"    description: {desc}\n"
                front_matter += f"    example: \"{example}\"\n"
                front_matter += f"    required: {str(required).lower()}\n"
        
        # Add similarity tags
        front_matter += f"similarity_tags: {tags}\n"
        front_matter += "---\n\n"
        
        # Replace variable patterns in content with {{key}} format
        template_content = content
        for var in variables:
            key = var['key']
            # Convert various patterns to {{key}} format
            patterns = [
                (f"[{key}]", f"{{{{{key}}}}}"),
                (f"_{key}_", f"{{{{{key}}}}}"),
                (f"<{key}>", f"{{{{{key}}}}}"),
                # Also handle already correct patterns
            ]
            for old, new in patterns:
                template_content = template_content.replace(old, new)
        
        return front_matter + template_content
    
    @staticmethod
    def generate_docx_from_markdown(markdown_text: str, output_path: str) -> None:
        """Convert markdown text to a DOCX file."""
        try:
            doc = Document()
            
            # Simple markdown to DOCX conversion
            lines = markdown_text.split('\n')
            in_front_matter = False
            
            for line in lines:
                line = line.strip()
                
                # Skip front-matter
                if line == '---':
                    in_front_matter = not in_front_matter
                    continue
                
                if in_front_matter:
                    continue
                
                if not line:
                    continue
                
                # Headers
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                # List items
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line[0].isdigit() and line[1:3] in ['. ', ') ']:
                    doc.add_paragraph(line[3:], style='List Number')
                # Regular paragraphs
                else:
                    # Handle bold/italic markdown
                    paragraph = doc.add_paragraph()
                    text = line
                    # Simple bold replacement
                    text = text.replace('**', '')
                    text = text.replace('__', '')
                    paragraph.add_run(text)
            
            doc.save(output_path)
            logger.info(f"DOCX generated successfully: {output_path}")
        
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    @staticmethod
    def export_variables_to_csv(variables: List[dict]) -> str:
        """Export template variables to CSV format."""
        import csv
        from io import StringIO
        
        output = StringIO()
        if not variables:
            return ""
        
        fieldnames = ['key', 'label', 'description', 'example', 'data_type', 'is_required', 'default_value']
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        
        writer.writeheader()
        for var in variables:
            writer.writerow({
                'key': var.get('key', ''),
                'label': var.get('label', var.get('key', '')),
                'description': var.get('description', ''),
                'example': var.get('example', ''),
                'data_type': var.get('data_type', 'text'),
                'is_required': var.get('is_required', True),
                'default_value': var.get('default_value', '')
            })
        
        return output.getvalue()

document_service = DocumentService()
